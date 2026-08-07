"""Session file attachments: raw uploads + extracted text for Lain.

Files live under ``data/ai-sessions/<project_id>/<session_id>/attachments/``:
- ``<id>.<ext>``    the original uploaded bytes
- ``<id>.txt``      extracted plain text (empty with an ``error`` note if parsing fails)
- ``metadata.json`` the session's attachment list

Attachments are never inside the project tree, so they can't collide with
user folders or leak into exports.
"""
from __future__ import annotations

import json
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from app import config
from app.ai import sessions

ALLOWED_EXTS = {".pdf", ".docx", ".txt", ".md"}
MAX_FILE_BYTES = 25 * 1024 * 1024
MAX_ATTACHMENTS = 20
MAX_EXTRACT_CHARS = 2_000_000


class AttachmentError(Exception):
    pass


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _dir(project_id: str, session_id: str) -> Path:
    return sessions.session_dir(project_id, session_id) / "attachments"


def _metadata_path(project_id: str, session_id: str) -> Path:
    return _dir(project_id, session_id) / "metadata.json"


def _load_metadata(project_id: str, session_id: str) -> list[dict[str, Any]]:
    path = _metadata_path(project_id, session_id)
    if not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return []
    return data if isinstance(data, list) else []


def _save_metadata(project_id: str, session_id: str, items: list[dict[str, Any]]) -> None:
    _dir(project_id, session_id).mkdir(parents=True, exist_ok=True)
    config._write_atomic(
        _metadata_path(project_id, session_id),
        json.dumps(items, ensure_ascii=False, indent=2),
    )


def _safe_ext(filename: str) -> str:
    ext = Path(filename or "").suffix.lower()
    if ext not in ALLOWED_EXTS:
        raise AttachmentError(
            f"Unsupported file type '{ext or '(none)'}' — allowed: pdf, docx, txt, md"
        )
    return ext


def list_attachments(project_id: str, session_id: str) -> list[dict[str, Any]]:
    items = _load_metadata(project_id, session_id)
    for item in items:
        item["tokens"] = max(1, int(item.get("chars") or 0) // 4)
    return items


def _extract_text(path: Path, ext: str) -> str:
    text = ""
    if ext in (".txt", ".md"):
        raw = path.read_bytes()
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("utf-8", errors="replace")
    elif ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    elif ext == ".docx":
        from docx import Document

        doc = Document(str(path))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text:
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
    return (text or "").strip()


def add(project_id: str, session_id: str, filename: str, raw: bytes) -> dict[str, Any]:
    ext = _safe_ext(filename)
    if len(raw) > MAX_FILE_BYTES:
        raise AttachmentError(f"File exceeds the 25 MB limit")
    items = _load_metadata(project_id, session_id)
    if len(items) >= MAX_ATTACHMENTS:
        raise AttachmentError(
            f"This session already has {MAX_ATTACHMENTS} attachments — remove one first"
        )
    aid = uuid.uuid4().hex
    d = _dir(project_id, session_id)
    d.mkdir(parents=True, exist_ok=True)
    raw_path = d / f"{aid}{ext}"
    raw_path.write_bytes(raw)
    error = None
    try:
        text = _extract_text(raw_path, ext)
    except Exception as exc:  # noqa: BLE001 — extraction is best-effort
        text = ""
        error = f"Could not extract text: {exc}"
    text = text[:MAX_EXTRACT_CHARS]
    config._write_atomic(d / f"{aid}.txt", text)
    item = {
        "id": aid,
        "name": Path(filename or "attachment").name,
        "ext": ext,
        "size": len(raw),
        "chars": len(text),
        "error": error,
        "createdAt": _now(),
    }
    items.append(item)
    _save_metadata(project_id, session_id, items)
    return item


def remove(project_id: str, session_id: str, attachment_id: str) -> bool:
    items = _load_metadata(project_id, session_id)
    found = False
    keep = []
    for item in items:
        if item.get("id") == attachment_id:
            found = True
            d = _dir(project_id, session_id)
            for path in d.glob(f"{attachment_id}.*"):
                try:
                    path.unlink()
                except OSError:
                    pass
        else:
            keep.append(item)
    if not found:
        return False
    _save_metadata(project_id, session_id, keep)
    return True


def get_text(project_id: str, session_id: str, attachment_id: str) -> str:
    path = _dir(project_id, session_id) / f"{attachment_id}.txt"
    if not path.exists():
        raise AttachmentError("Attachment not found")
    return path.read_text(encoding="utf-8")
