"""Shared export utilities: markdown parsing, DOCX/HTML generation."""
from __future__ import annotations

import re
import yaml
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import markdown as md_lib

from app import config
from app.services import documents as documents_service


_FRONTMATTER_RE = re.compile(r"^---\s*\n(.*?)\n---\s*\n?", re.DOTALL)
_WIKILINK_RE = re.compile(r"\[\[([^\]]+)\]\]")


def parse_frontmatter(text: str) -> tuple[dict[str, Any], str]:
    m = _FRONTMATTER_RE.match(text)
    if m:
        try:
            meta = yaml.safe_load(m.group(1)) or {}
        except yaml.YAMLError:
            meta = {}
        body = text[m.end():]
        return dict(meta), body
    return {}, text


def md_to_html(body: str) -> str:
    body = _WIKILINK_RE.sub(r"\1", body)
    return md_lib.markdown(
        body,
        extensions=["extra", "codehilite", "sane_lists"],
        output_format="html5",
    )


def build_export_html(project_title: str, documents: list[tuple[str, str, str]]) -> str:
    parts = []
    parts.append(f"<h1 class='project-title'>{_esc(project_title)}</h1>")

    current_folder = None
    for folder_name, doc_title, body in documents:
        if folder_name != current_folder:
            if current_folder is not None:
                parts.append('<div class="page-break"></div>')
            parts.append(f"<h2 class='folder-heading'>{_esc(folder_name)}</h2>")
            current_folder = folder_name
        parts.append(f"<h3 class='doc-title'>{_esc(doc_title)}</h3>")
        parts.append(md_to_html(body))

    return """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{title}</title>
<style>
  @page {{ size: A4; margin: 2cm 2.5cm; @bottom-center {{ content: counter(page); font-size: 9pt; color: #888; }} }}
  body {{ font-family: "Georgia", "Times New Roman", serif; font-size: 12pt; line-height: 1.6; color: #1a1a1a; }}
  h1.project-title {{ font-size: 22pt; text-align: center; margin-bottom: 1.5cm; }}
  h2.folder-heading {{ font-size: 16pt; margin-top: 2em; border-bottom: 1px solid #ccc; padding-bottom: 0.3em; }}
  h3.doc-title {{ font-size: 14pt; margin-top: 1.5em; }}
  h4 {{ font-size: 13pt; }}
  h5, h6 {{ font-size: 12pt; }}
  p {{ margin: 0.5em 0; text-align: justify; }}
  blockquote {{ margin: 1em 2em; font-style: italic; color: #444; border-left: 3px solid #bbb; padding-left: 1em; }}
  code {{ font-family: "Courier New", monospace; font-size: 10pt; background: #f5f5f5; padding: 1px 4px; }}
  pre {{ background: #f5f5f5; border: 1px solid #ddd; padding: 0.8em; font-size: 10pt; overflow-x: auto; }}
  pre code {{ background: none; padding: 0; }}
  hr {{ border: none; border-top: 1px solid #ccc; margin: 1.5em 0; }}
  ul, ol {{ margin: 0.5em 0; padding-left: 2em; }}
  li {{ margin: 0.2em 0; }}
  .page-break {{ page-break-before: always; }}
</style>
</head>
<body>
{body}
</body>
</html>""".format(title=_esc(project_title), body="\n".join(parts))


def _esc(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


# ---------------------------------------------------------------------------
# Document collection
# ---------------------------------------------------------------------------

def _iter_md_files(project_folder: Path, folder_ids: list[str] | None) -> list[tuple[Path, str]]:
    selected = None
    root_docs = False
    if folder_ids:
        selected = set()
        for fid in folder_ids:
            fid = fid.strip()
            if fid == ".":
                root_docs = True
            elif fid:
                selected.add(fid.rstrip("/").replace("\\", "/"))

    result: list[tuple[Path, str]] = []
    for path in sorted(project_folder.rglob("*.md"), key=lambda p: _sort_path(p, project_folder)):
        rel = path.relative_to(project_folder)
        if ".reorder-tmp" in rel.parts or rel.name.startswith("."):
            continue
        if selected is not None:
            parents = list(rel.parent.parts) if rel.parent != Path(".") else []
            if not _path_matches(parents, selected) and not (root_docs and not parents):
                continue
        folder_name = "/".join(documents_service._display_name(p) for p in rel.parent.parts) if rel.parent != Path(".") else ""
        result.append((path, folder_name))
    return result


def _path_matches(parents: list[str], selected: set[str]) -> bool:
    if not parents:
        return False
    for i in range(1, len(parents) + 1):
        prefix = "/".join(parents[:i])
        if prefix in selected:
            return True
    return False


def _sort_path(path: Path, project_folder: Path) -> tuple:
    parts = path.relative_to(project_folder).parts
    return tuple(documents_service._sort_key(p) for p in parts)


def collect_documents(project_id: str, folder_ids: list[str] | None) -> list[tuple[str, str, str]]:
    """Return list of (folder_display_name, doc_title, body) sorted by order."""
    from app.services.projects import project_dir, _safe_id
    pid = _safe_id(project_id)
    folder = project_dir(pid)
    if not folder.exists():
        raise FileNotFoundError(f"Project '{project_id}' not found")

    result: list[tuple[str, str, str]] = []
    for path, folder_name in _iter_md_files(folder, folder_ids):
        raw = path.read_text(encoding="utf-8", errors="replace")
        meta, body = parse_frontmatter(raw)
        title = meta.get("title") or documents_service._display_name(path.stem).replace("-", " ").title()
        title = str(title) if title is not None else "Untitled"
        result.append((folder_name, title, body.strip()))
    return result


# ---------------------------------------------------------------------------
# HTML -> DOCX converter
# ---------------------------------------------------------------------------

class _DocxBuilder(HTMLParser):
    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self._para = None
        self._run = None
        self._stack: list[dict] = []
        self._list_depth = 0
        self._ol_counters: list[int] = []

    def _push(self, tag: str, attrs):
        fmt = {}
        for k, v in attrs:
            if k == "class" and "codehilite" in (v or ""):
                fmt["code_block"] = True
        self._stack.append({"tag": tag, **fmt})

    def _pop(self):
        if self._stack:
            self._stack.pop()

    def _top_tag(self) -> str | None:
        return self._stack[-1]["tag"] if self._stack else None

    def handle_starttag(self, tag, attrs):
        attrs_dict = {k: v for k, v in attrs}
        tt = self._top_tag()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._para = None
            self._push(tag, attrs)
        elif tag == "p":
            self._para = None
            self._push(tag, attrs)
        elif tag == "strong" or tag == "b":
            self._push(tag, attrs)
        elif tag == "em" or tag == "i":
            self._push(tag, attrs)
        elif tag == "code":
            if tt == "pre":
                pass
            else:
                self._push(tag, attrs)
        elif tag == "a":
            self._push(tag, attrs)
        elif tag == "blockquote":
            self._para = None
            self._push(tag, attrs)
        elif tag == "pre":
            self._para = None
            self._push(tag, attrs)
            self._code_lines = []
        elif tag == "ul":
            self._list_depth += 1
            self._push(tag, attrs)
        elif tag == "ol":
            self._list_depth += 1
            self._ol_counters.append(0)
            self._push(tag, attrs)
        elif tag == "li":
            self._push(tag, attrs)
        elif tag == "hr":
            self.doc.add_paragraph("_" * 60)
            p = self.doc.paragraphs[-1]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            self._para = None
        elif tag == "br":
            if self._para is not None:
                self._run = self._para.add_run("\n")

    def handle_endtag(self, tag):
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._pop()
            self._para = None
        elif tag == "p":
            self._pop()
            self._para = None
        elif tag in ("strong", "b", "em", "i", "code", "a"):
            self._pop()
        elif tag == "blockquote":
            self._pop()
            self._para = None
        elif tag == "pre":
            self._pop()
            text = "".join(getattr(self, "_code_lines", []))
            p = self.doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            pf = p.paragraph_format
            pf.left_indent = Inches(0.5)
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            self._para = None
        elif tag == "ul":
            self._list_depth -= 1
            self._pop()
        elif tag == "ol":
            self._list_depth -= 1
            if self._ol_counters:
                self._ol_counters.pop()
            self._pop()

    def handle_data(self, data):
        tt = self._top_tag()

        if tt is None or tt == "body":
            return

        if tt == "pre":
            self._code_lines.append(data)
            return

        if self._para is None:
            if tt in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level = int(tt[1])
                self._para = self.doc.add_heading("", level=level)
                self._run = self._para.add_run(data)
            elif tt == "blockquote":
                self._para = self.doc.add_paragraph()
                pf = self._para.paragraph_format
                pf.left_indent = Inches(0.5)
                self._run = self._para.add_run(data)
                self._run.font.italic = True
            elif tt == "p":
                self._para = self.doc.add_paragraph()
                self._run = self._para.add_run(data)
            elif tt == "li":
                if self._list_depth > 0 and self._ol_counters and len(self._ol_counters) >= self._list_depth:
                    self._ol_counters[self._list_depth - 1] += 1
                    num = self._ol_counters[self._list_depth - 1]
                    self._para = self.doc.add_paragraph(style="List Number")
                else:
                    self._para = self.doc.add_paragraph(style="List Bullet")
                self._run = self._para.add_run(data)
            else:
                return
        else:
            self._run = self._get_run()
            self._run.add_text(data)

    def _get_run(self):
        if self._para is None:
            self._para = self.doc.add_paragraph()
        run = self._para.add_run()
        attrs = self._cumulative_fmt()
        if "strong" in attrs:
            run.bold = True
        if "em" in attrs:
            run.italic = True
        if "code" in attrs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        if "a" in attrs:
            run.underline = True
            run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
        return run

    def _cumulative_fmt(self) -> set:
        fmt: set = set()
        for item in self._stack:
            tag = item["tag"]
            if tag == "strong" or tag == "b":
                fmt.add("strong")
            elif tag == "em" or tag == "i":
                fmt.add("em")
            elif tag == "code":
                fmt.add("code")
            elif tag == "a":
                fmt.add("a")
        return fmt


def html_to_docx(doc: Document, html: str, title: str | None = None) -> None:
    if title:
        doc.add_heading(str(title), level=1)
    parser = _DocxBuilder(doc)
    parser.feed(html)
    parser.close()
