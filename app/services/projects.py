"""Project management: one folder per project on disk."""
from __future__ import annotations

import json
import re
import shutil
import unicodedata
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from app import config
from app.services import documents as documents_service


def _slugify(name: str) -> str:
    """Return a filesystem-safe slug from a display name."""
    normalized = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode("ascii")
    normalized = re.sub(r"[^A-Za-z0-9]+", "-", normalized).strip("-").lower()
    return normalized or "project"


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def project_dir(project_id: str) -> Path:
    return config.DATA_DIR / project_id


def _safe_id(project_id: str) -> str:
    if not re.fullmatch(r"[A-Za-z0-9._-]+", project_id):
        raise ValueError("Invalid project id")
    return project_id


def _read_meta(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        data = {}
    return {
        "title": data.get("title", path.parent.name),
        "goal": data.get("goal", {"wordsPerDay": 500, "enabled": False}),
        "createdAt": data.get("createdAt", ""),
        "updatedAt": data.get("updatedAt", ""),
        **data,
    }


def _write_meta(project: Path, meta: dict[str, Any]) -> None:
    (project / config.PROJECT_META_FILENAME).write_text(
        json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def list_projects() -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    if not config.DATA_DIR.exists():
        return results
    for folder in sorted(config.DATA_DIR.iterdir(), key=lambda p: p.name.lower()):
        meta_path = folder / config.PROJECT_META_FILENAME
        if not folder.is_dir() or not meta_path.exists():
            continue
        meta = _read_meta(meta_path)
        stats = documents_service.project_word_stats(folder.name)
        results.append(
            {
                "id": folder.name,
                "title": meta["title"],
                "goal": meta["goal"],
                "createdAt": meta["createdAt"],
                "updatedAt": meta["updatedAt"],
                "words": stats["words"],
                "documents": stats["documents"],
            }
        )
    return results


def get_project(project_id: str) -> dict[str, Any]:
    pid = _safe_id(project_id)
    folder = project_dir(pid)
    meta_path = folder / config.PROJECT_META_FILENAME
    if not meta_path.exists():
        raise FileNotFoundError(f"Project '{project_id}' not found")
    meta = _read_meta(meta_path)
    stats = documents_service.project_word_stats(pid)
    return {
        "id": pid,
        "title": meta["title"],
        "goal": meta["goal"],
        "createdAt": meta["createdAt"],
        "updatedAt": meta["updatedAt"],
        "words": stats["words"],
        "documents": stats["documents"],
    }


def create_project(name: str) -> dict[str, Any]:
    slug = _slugify(name)
    folder = project_dir(slug)
    if folder.exists():
        raise ValueError(f"Project '{name}' already exists")
    folder.mkdir(parents=True, exist_ok=True)
    (folder / config.STATS_DIRNAME).mkdir(exist_ok=True)
    (folder / config.WIKI_DIRNAME).mkdir(exist_ok=True)
    now = _now()
    meta = {
        "title": name.strip() or slug,
        "goal": {"wordsPerDay": 500, "enabled": False},
        "createdAt": now,
        "updatedAt": now,
    }
    _write_meta(folder, meta)
    return get_project(slug)


def rename_project(project_id: str, new_name: str) -> dict[str, Any]:
    pid = _safe_id(project_id)
    folder = project_dir(pid)
    if not folder.exists():
        raise FileNotFoundError(f"Project '{project_id}' not found")
    new_slug = _slugify(new_name)
    new_folder = project_dir(new_slug)
    if new_folder.exists() and new_folder != folder:
        raise ValueError(f"Project '{new_name}' already exists")
    meta = _read_meta(folder / config.PROJECT_META_FILENAME)
    meta["title"] = new_name.strip() or meta["title"]
    meta["updatedAt"] = _now()
    if new_folder != folder:
        folder.rename(new_folder)
        folder = new_folder
    _write_meta(folder, meta)
    return get_project(new_folder.name)


def set_goal(project_id: str, words_per_day: int, enabled: bool) -> dict[str, Any]:
    pid = _safe_id(project_id)
    folder = project_dir(pid)
    meta_path = folder / config.PROJECT_META_FILENAME
    if not meta_path.exists():
        raise FileNotFoundError(f"Project '{project_id}' not found")
    meta = _read_meta(meta_path)
    meta["goal"] = {
        "wordsPerDay": max(1, int(words_per_day)),
        "enabled": bool(enabled),
    }
    meta["updatedAt"] = _now()
    _write_meta(folder, meta)
    return get_project(pid)


def delete_project(project_id: str) -> None:
    pid = _safe_id(project_id)
    folder = project_dir(pid)
    if not folder.exists():
        raise FileNotFoundError(f"Project '{project_id}' not found")
    shutil.rmtree(folder)


def get_document_tree(project_id: str, scope: str = "write") -> dict[str, Any]:
    return documents_service.get_tree(project_id, scope=scope)


def export_zip(project_id: str, folder_ids: list[str] | None = None) -> bytes:
    """Return the project folder (Markdown files + metadata) as a zip archive.

    The hidden ``NN-`` order prefixes are stripped from file/folder names in
    the archive; if two entries would collide after stripping, the later one
    gets a ``-2``/``-3``… suffix so nothing is overwritten.

    If *folder_ids* is a non-empty list only files inside those folders (or
    project-level files when ``"."`` is included) are included.  Core metadata
    (``project.json``, ``stats/``) is always included.
    """
    import io
    import zipfile

    from app.services.export import _iter_md_files

    pid = _safe_id(project_id)
    folder = project_dir(pid)
    if not folder.exists():
        raise FileNotFoundError(f"Project '{project_id}' not found")

    if folder_ids is not None and len(folder_ids) == 0:
        folder_ids = None

    included: set[Path] | None = None
    if folder_ids is not None:
        included = set()
        for path, _ in _iter_md_files(folder, folder_ids):
            included.add(path)
        included.add(folder / config.PROJECT_META_FILENAME)
        for p in (folder / config.STATS_DIRNAME).rglob("*"):
            if p.is_file() and ".reorder-tmp" not in p.parts:
                included.add(p)

    buffer = io.BytesIO()
    used: set[str] = set()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for path in sorted(folder.rglob("*")):
            if not path.is_file() or ".reorder-tmp" in path.parts:
                continue
            if included is not None and path not in included:
                continue
            parts = path.relative_to(folder).as_posix().split("/")
            arcname = "/".join(documents_service._display_name(p) for p in parts)
            if arcname in used:
                dot = arcname.rfind(".")
                base, ext = (arcname[:dot], arcname[dot:]) if dot > 0 else (arcname, "")
                n = 2
                while f"{base}-{n}{ext}" in used:
                    n += 1
                arcname = f"{base}-{n}{ext}"
            used.add(arcname)
            zf.write(path, arcname=arcname)
    buffer.seek(0)
    return buffer.getvalue()


def export_docx(project_id: str, folder_ids: list[str] | None = None) -> bytes:
    """Build a single .docx file from selected markdown documents."""
    import io

    from docx import Document
    from docx.shared import Pt

    from app.services.export import collect_documents, md_to_html, html_to_docx

    pid = _safe_id(project_id)
    meta = _read_meta(project_dir(pid) / config.PROJECT_META_FILENAME)
    project_title = meta.get("title", pid)

    docs = collect_documents(project_id, folder_ids)

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Georgia"
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    doc.add_heading(project_title, level=0)

    current_folder = None
    for folder_name, doc_title, body in docs:
        if folder_name != current_folder:
            current_folder = folder_name
            label = folder_name or "Top-level"
            doc.add_heading(label, level=2)
        html = md_to_html(body)
        html_to_docx(doc, html, title=doc_title)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_pdf(project_id: str, folder_ids: list[str] | None = None) -> bytes:
    """Build a single .pdf file from selected markdown documents."""
    import io

    from fpdf import FPDF
    from fpdf.fonts import FontFace

    from app.services.export import collect_documents, md_to_html

    pid = _safe_id(project_id)
    meta = _read_meta(project_dir(pid) / config.PROJECT_META_FILENAME)
    project_title = meta.get("title", pid)

    docs = collect_documents(project_id, folder_ids)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    pdf.add_font("Serif", "", r"C:\Windows\Fonts\georgia.ttf", uni=True)
    pdf.add_font("Serif", "B", r"C:\Windows\Fonts\georgiab.ttf", uni=True)
    pdf.add_font("Serif", "I", r"C:\Windows\Fonts\georgiai.ttf", uni=True)
    pdf.add_font("Serif", "BI", r"C:\Windows\Fonts\georgiaz.ttf", uni=True)
    pdf.add_font("Mono", "", r"C:\Windows\Fonts\cour.ttf", uni=True)

    pdf.set_font("Serif", "B", 22)
    pdf.cell(0, 14, project_title, new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    current_folder = None
    for folder_name, doc_title, body in docs:
        if folder_name != current_folder:
            current_folder = folder_name
            label = folder_name or "Top-level"
            if pdf.get_y() > 240:
                pdf.add_page()
            pdf.set_font("Serif", "B", 14)
            pdf.set_draw_color(200, 200, 200)
            pdf.set_line_width(0.4)
            y = pdf.get_y()
            pdf.cell(0, 8, label, new_x="LMARGIN", new_y="NEXT")
            pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + pdf.w - 2 * pdf.l_margin, pdf.get_y())
            pdf.ln(4)

        pdf.set_font("Serif", "B", 12)
        pdf.cell(0, 7, str(doc_title), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        html = md_to_html(body)
        pdf.set_font("Serif", "", 11)
        pdf.write_html(html, tag_styles={
            "h1": FontFace(family="Serif", emphasis="B", size_pt=16),
            "h2": FontFace(family="Serif", emphasis="B", size_pt=14),
            "h3": FontFace(family="Serif", emphasis="B", size_pt=12),
            "h4": FontFace(family="Serif", emphasis="B", size_pt=11),
            "code": FontFace(family="Mono", size_pt=9),
            "pre": FontFace(family="Mono", size_pt=9),
            "blockquote": FontFace(family="Serif", emphasis="I", color=(100, 100, 100)),
        })
        pdf.ln(6)

    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_epub(project_id: str, folder_ids: list[str] | None = None) -> bytes:
    """Build a single .epub file from selected markdown documents."""
    import io
    import uuid

    from ebooklib import epub

    from app.services.export import collect_documents, md_to_html

    pid = _safe_id(project_id)
    meta = _read_meta(project_dir(pid) / config.PROJECT_META_FILENAME)
    project_title = meta.get("title", pid)

    docs = collect_documents(project_id, folder_ids)

    book = epub.EpubBook()
    book.set_identifier(str(uuid.uuid4()))
    book.set_title(project_title)
    book.set_language("en")
    book.add_author(project_title)

    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    spine = ["nav"]
    toc = []
    current_folder = None

    css = epub.EpubItem(
        uid="style",
        file_name="style/default.css",
        media_type="text/css",
        content=b"""body { font-family: Georgia, serif; line-height: 1.6; margin: 1em; }
h1 { text-align: center; font-size: 1.8em; margin: 1em 0; }
h2 { font-size: 1.3em; margin: 1.2em 0 0.4em; border-bottom: 1px solid #ccc; padding-bottom: 0.2em; }
h3 { font-size: 1.15em; margin: 1em 0 0.3em; }
h4 { font-size: 1.05em; }
blockquote { margin: 1em 2em; font-style: italic; color: #444; border-left: 3px solid #bbb; padding-left: 1em; }
pre, code { font-family: "Courier New", monospace; font-size: 0.9em; }
pre { background: #f5f5f5; padding: 0.6em; }
p { margin: 0.5em 0; }""",
    )
    book.add_item(css)

    for folder_name, doc_title, body in docs:
        if folder_name != current_folder:
            current_folder = folder_name
            label = folder_name or "Top-level"
            sec = epub.EpubHtml(
                title=label,
                file_name=f"sec_{_slugify_short(label)}_{len(spine)}.xhtml",
                lang="en",
            )
            sec.content = f"<h2>{label}</h2>".encode("utf-8")
            sec.add_item(css)
            book.add_item(sec)
            spine.append(sec)
            toc.append(epub.Link(sec.file_name, label, f"toc_{len(spine)}"))

        html_body = md_to_html(body)
        full_html = f"<h3>{doc_title}</h3>\n{html_body}"
        ch = epub.EpubHtml(
            title=doc_title,
            file_name=f"ch_{_slugify_short(doc_title)}_{len(spine)}.xhtml",
            lang="en",
        )
        ch.content = full_html.encode("utf-8")
        ch.add_item(css)
        book.add_item(ch)
        spine.append(ch)
        toc.append(epub.Link(ch.file_name, doc_title, f"toc_{len(spine)}"))

    book.toc = toc
    book.spine = spine

    buffer = io.BytesIO()
    epub.write_epub(buffer, book)
    buffer.seek(0)
    return buffer.getvalue()


def _slugify_short(name: str) -> str:
    import unicodedata
    n = unicodedata.normalize("NFKD", name).encode("ascii", "ignore").decode("ascii")
    n = re.sub(r"[^A-Za-z0-9]+", "-", n).strip("-").lower()
    return n[:30] or "x"
